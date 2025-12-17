"""
Document processor for extracting text from various file formats.
"""
import os
from typing import BinaryIO
from PyPDF2 import PdfReader
from docx import Document
import tempfile


class DocumentProcessor:
    """Process various document types and extract text content."""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}

    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Get the file extension from filename."""
        return os.path.splitext(filename)[1].lower()

    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if file type is supported."""
        ext = DocumentProcessor.get_file_extension(filename)
        return ext in DocumentProcessor.SUPPORTED_EXTENSIONS

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file.

        Args:
            file_content: PDF file content as bytes

        Returns:
            Extracted text content
        """
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_file_path = tmp_file.name

        try:
            reader = PdfReader(tmp_file_path)
            text_content = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)

            return "\n\n".join(text_content)
        finally:
            os.unlink(tmp_file_path)

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Extracted text content
        """
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_file_path = tmp_file.name

        try:
            doc = Document(tmp_file_path)
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_content.append(" | ".join(row_text))

            return "\n\n".join(text_content)
        finally:
            os.unlink(tmp_file_path)

    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """
        Extract text from TXT file.

        Args:
            file_content: TXT file content as bytes

        Returns:
            Text content
        """
        try:
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            return file_content.decode("latin-1")

    @classmethod
    def extract_text(cls, filename: str, file_content: bytes) -> str:
        """
        Extract text from a document based on its type.

        Args:
            filename: Name of the file (to determine type)
            file_content: File content as bytes

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported
        """
        ext = cls.get_file_extension(filename)

        if ext == ".pdf":
            return cls.extract_text_from_pdf(file_content)
        elif ext == ".docx":
            return cls.extract_text_from_docx(file_content)
        elif ext == ".txt":
            return cls.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
